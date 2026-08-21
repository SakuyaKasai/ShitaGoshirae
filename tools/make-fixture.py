#!/usr/bin/env python3
"""
テスト用の「素の原稿」docx を作る。

意図的に、実際の原稿で観測された問題を全部盛り込む:
  - 句点「。」読点「、」の混在（W01/W02。実データでは句点93箇所）
  - 見出し番号の記法ゆれ（全角「１．」と半角「2-1.」— W03）
  - 見出し番号の欠番（W04）
  - Unicode のベタ書き数式（W07）
  - 単位の前にスペースがない（W15）
  - 参考文献に未引用のものがある（W05）
  - Wordのスタイルを一切使っていない（＝先方の原稿の実態）
  - 画像・表・箇条書きを含む
"""
import sys
from docx import Document
from docx.shared import Pt, Mm
from PIL import Image, ImageDraw

out = sys.argv[1] if len(sys.argv) > 1 else 'fixture-source.docx'

# --- ダミー画像を2枚作る（片方は低解像度にして W12 を踏ませる）---
def make_png(path, w, h, label, color):
    img = Image.new('RGB', (w, h), 'white')
    d = ImageDraw.Draw(img)
    d.rectangle([4, 4, w - 5, h - 5], outline=color, width=3)
    d.text((14, h // 2), label, fill=color)
    img.save(path)

make_png('fig1.png', 900, 640, 'Figure 1: load path', (30, 60, 140))
make_png('fig2.png', 220, 160, 'Figure 2: low-res', (140, 40, 40))   # → W12

doc = Document()

def para(text, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

# ===== 本文（章立てあり。スタイルは一切使わない）=====
para('１．はじめに')                       # 全角記法
para('着物の着付けにおける腰紐の荷重流路を数値解析するため、構造の変換手法を検討した。'
     '本研究では静力学的なモデルを用いる。先行研究1)では十分に扱われていない。')
para('サンプリング周波数は32kHzとし、質量は5.5kgであった。2024年に実施した。')

para('２．方法')
para('2-1. 対象')                          # 半角記法（W03: 記法ゆれ）
para('被験者は成人女性10名とした。平均年齢は42.3歳、身長は158.4cmであった。')
para('ΣᵢFᵢ＝0')                            # W07
para('M(r)＝(rL－r)×F')                     # W07
para('2-3. 解析条件')                      # W04: 2-2 が欠番
para('境界条件を以下のように設定した。')
para('・腰紐の張力は一定とする')
para('・摩擦係数は0.3とする')

# 表（w:tbl として運ばれるべきもの）
t = doc.add_table(rows=3, cols=3)
t.style = 'Table Grid'
cells = [['条件', '値', '単位'], ['張力', '12.5', 'N'], ['摩擦係数', '0.3', '—']]
for i, row in enumerate(cells):
    for j, v in enumerate(row):
        t.cell(i, j).text = v
para('表1．解析条件の一覧')

para('３．結果')
doc.add_picture('fig1.png', width=Mm(70))
para('図1．荷重流路の模式図')
doc.add_picture('fig2.png', width=Mm(70))   # 低解像度 → W12
para('図2．低解像度の例')
para('解析の結果、荷重は腰紐に沿って伝達されることが確認された。'
     '最大応力は12.5MPaであった。文献2)の報告と整合する。')

para('4. 考察')
para('本手法により、着付け構造を数値解析条件へ変換できることが示された。'
     '今後は動的な条件への拡張が課題である5)。')   # W06: 文献5) は存在しない

para('文献')
for ref in [
    '1) 大須賀美恵子, 青木和夫. 座談会. 人間工学. 2014.',
    '2) Dul, J.; Bruder, R. Ergonomics. 2012.',
    '3) French, J. C.; Chapin, A. Textile Research. 2019.',   # W05: 未引用
]:
    para(ref)

doc.save(out)
print(f'✅ {out} を生成しました')
