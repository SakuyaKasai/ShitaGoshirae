#!/usr/bin/env python3
"""
書誌情報の自動抽出を試すための原稿を作る。

make-fixture.py が「本文の問題」を再現するのに対し、
こちらは「冒頭の表題・著者・抄録・キーワード・所属」を持った、
実際に投稿される形に近い原稿を作る。

テンプレの見本を上書きして書いた原稿を想定しているので、
「■技術報告■」「, by」「（キーワード：」といった
テンプレ由来のアンカー文字列が残っている。
抽出アルゴリズムはこれらを手がかりにする。
"""
import sys
from docx import Document
from docx.shared import Pt

out = sys.argv[1] if len(sys.argv) > 1 else 'fixture-full.docx'
doc = Document()


def p(text):
    q = doc.add_paragraph()
    r = q.add_run(text)
    r.font.size = Pt(10.5)
    return q


# ===== 冒頭（セクション1の材料）=====
p('■技術報告■\t\t<Technical report> A Conversion Method from Women\'s Kimono '
  'Waist-Cord Dressing Structures to Numerical Analysis Conditions, '
  'by Yuri AOKI & Miho KASAI')
p('女性腰紐着付け構造の数値解析条件への変換手法1')          # 末尾の 1 は脚注番号
p('青木　友里2，笠井　美穂3')                                # 末尾の数字は所属番号
p('This study proposes a conversion method from the dressing structures of '
  'women\'s kimono waist-cords into numerical analysis conditions. ' * 6)
p('本研究では，女性の着物着付けにおける腰紐の構造を数値解析条件へ変換する手法を提案する．' * 5)
p('（キーワード：　着物，着付け，腰紐，荷重流路，静力学）')
p('2 人間工学大学人間工学学部 / School of Ergonomics, Ningenkougaku University')
p('3 株式会社アーゴノミクス / Ergonomics Co. Ltd.')

# ===== 本文（ここから先が bodyStart）=====
p('1. はじめに')
p('着物の着付けにおける腰紐の荷重流路を数値解析するため、構造の変換手法を検討した。'
  '先行研究1)では十分に扱われていない。')
p('2. 方法')
p('2-1. 対象')
p('被験者は成人女性10名とした。サンプリング周波数は32kHzであった。')
p('文献')
p('1) 大須賀美恵子. 座談会. 人間工学. 2014.')

doc.save(out)
print(f'✅ {out} を生成しました')
